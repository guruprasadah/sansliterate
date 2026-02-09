#!/usr/bin/env python3
"""
Transliterate Sanskrit (Devanagari) text in a DOCX file to Tamil script.

Workflow:
- Treat the input DOCX as read-only.
- Create a copy (or a separate output path).
- Traverse paragraphs and tables in the copied document.
- Within each run, detect spans that contain Sanskrit (Devanagari) characters.
- Transliterate only those Sanskrit spans to Tamil, leaving other text and formatting intact.

Dependencies:
- python-docx (pip install python-docx)
- indic-transliteration (pip install indic-transliteration)

Usage example:
    python scripts/transliterate_docx_sanskrit_to_tamil.py input.docx -o output_ta.docx
    python scripts/transliterate_docx_sanskrit_to_tamil.py input.docx --dry-run -v
    python scripts/transliterate_docx_sanskrit_to_tamil.py input.docx -o output_ta.docx --tamil-font Vijaya
"""

from __future__ import annotations

import argparse
import logging
import shutil
from pathlib import Path
from typing import Iterable, List, Sequence, Tuple

from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore

try:
    from docx import Document  # type: ignore
    from docx.document import Document as _DocxDocument  # type: ignore
    from docx.table import _Cell, Table  # type: ignore
    from docx.text.paragraph import Paragraph  # type: ignore
except ImportError:  # pragma: no cover - import-time dependency check
    Document = None  # type: ignore[assignment]
    _DocxDocument = None  # type: ignore[assignment]
    _Cell = None  # type: ignore[assignment]
    Table = None  # type: ignore[assignment]
    Paragraph = None  # type: ignore[assignment]

try:
    from indic_transliteration import sanscript  # type: ignore
    from indic_transliteration.sanscript import transliterate as indic_transliterate  # type: ignore
except ImportError:  # pragma: no cover - import-time dependency check
    sanscript = None  # type: ignore[assignment]
    indic_transliterate = None  # type: ignore[assignment]


LOGGER = logging.getLogger(__name__)

# Unicode range for Devanagari characters (U+0900–U+097F)
DEVANAGARI_START = 0x0900
DEVANAGARI_END = 0x097F


def is_sanskrit_char(ch: str) -> bool:
    """Return True if the character is in the Devanagari Unicode block."""
    if not ch:
        return False
    codepoint = ord(ch)
    return DEVANAGARI_START <= codepoint <= DEVANAGARI_END


Span = Tuple[str, str]  # (kind, text) where kind is "sanskrit" or "other"


def split_into_spans(text: str) -> List[Span]:
    """
    Split a string into spans of consecutive Sanskrit vs non-Sanskrit characters.

    A span is labeled "sanskrit" if its characters are in the Devanagari block,
    otherwise "other". Whitespace, punctuation, and digits are treated as "other"
    and thus left unchanged.
    """
    if not text:
        return []

    spans: List[Span] = []
    current_is_sanskrit = is_sanskrit_char(text[0])
    current_chars: List[str] = [text[0]]

    for ch in text[1:]:
        ch_is_sanskrit = is_sanskrit_char(ch)
        if ch_is_sanskrit == current_is_sanskrit:
            current_chars.append(ch)
        else:
            kind = "sanskrit" if current_is_sanskrit else "other"
            spans.append((kind, "".join(current_chars)))
            current_chars = [ch]
            current_is_sanskrit = ch_is_sanskrit

    kind = "sanskrit" if current_is_sanskrit else "other"
    spans.append((kind, "".join(current_chars)))
    return spans


def transliterate_sanskrit_to_tamil(text: str) -> str:
    """
    Transliterate Devanagari (Sanskrit) text to Tamil script.

    This function expects that `text` contains Devanagari characters. It uses the
    `indic-transliteration` library under the hood. If the library is not
    installed, the script will exit with an instructive error message.
    """
    if indic_transliterate is None or sanscript is None:
        LOGGER.error(
            "Missing dependency 'indic-transliteration'. "
            "Install it with: pip install indic-transliteration"
        )
        raise SystemExit(1)

    # The indic-transliteration library preserves punctuation and non-letter
    # characters, so we can pass the span directly.
    return indic_transliterate(text, sanscript.DEVANAGARI, sanscript.TAMIL)


def transliterate_run_text(text: str) -> str:
    """
    Transliterate only the Sanskrit (Devanagari) parts of a run's text to Tamil.

    Non-Sanskrit segments (including English, existing Tamil, punctuation, etc.)
    are left unchanged.
    """
    spans = split_into_spans(text)
    if not spans:
        return text

    # If there are no Sanskrit spans at all, avoid calling the transliterator.
    if not any(kind == "sanskrit" for kind, _ in spans):
        return text

    parts: List[str] = []
    for kind, segment in spans:
        if kind == "sanskrit":
            parts.append(transliterate_sanskrit_to_tamil(segment))
        else:
            parts.append(segment)
    return "".join(parts)


def should_apply_tamil_font(original_text: str) -> bool:
    """
    Decide whether it is safe to set the Tamil font on this run.

    We apply the Tamil font only when:
    - The run contained at least one Sanskrit (Devanagari) character, and
    - It does NOT contain any non-Sanskrit alphabetic characters
      (so mixed-script runs like Sanskrit+English keep their existing font).
    """

    has_sanskrit = False
    has_non_sanskrit_alpha = False

    for ch in original_text:
        if is_sanskrit_char(ch):
            has_sanskrit = True
        elif ch.isalpha():
            has_non_sanskrit_alpha = True

    return has_sanskrit and not has_non_sanskrit_alpha


def iter_paragraphs(parent: object) -> Iterable["Paragraph"]:
    """
    Yield all Paragraph objects from a document, including those inside tables.

    This intentionally traverses:
    - Document body paragraphs
    - Table cell paragraphs
    - Nested tables within cells (recursively)
    """
    if _DocxDocument is None or Table is None or _Cell is None or Paragraph is None:
        LOGGER.error(
            "Missing dependency 'python-docx'. Install it with: pip install python-docx"
        )
        raise SystemExit(1)

    if isinstance(parent, _DocxDocument):
        # Top-level document: body paragraphs and tables
        for paragraph in parent.paragraphs:
            yield paragraph
        for table in parent.tables:
            yield from iter_paragraphs(table)
    elif isinstance(parent, Table):
        for row in parent.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)
    elif isinstance(parent, _Cell):
        for paragraph in parent.paragraphs:
            yield paragraph
        for table in parent.tables:
            yield from iter_paragraphs(table)
    else:
        # Fallback: if an unexpected type, do nothing
        return


def process_document(
    doc: "Document",
    *,
    dry_run: bool = False,
    restrict_style: str | None = None,
    tamil_font: str | None = "Vijaya",
) -> Tuple[int, int]:
    """
    Process the document in-place, transliterating Sanskrit spans to Tamil.

    Args:
        doc: Loaded python-docx Document for the (copied) DOCX file.
        dry_run: If True, do not mutate the document; only count potential changes.
        restrict_style: Optional run style name. If provided, only runs whose
            style name exactly matches this string are considered for
            transliteration. This is useful when Sanskrit is always marked with
            a specific style in the document.
        tamil_font: Optional font name to apply to runs that are entirely
            Sanskrit (Devanagari) before transliteration. Defaults to "Vijaya".

    Returns:
        (total_runs, modified_runs)
    """
    total_runs = 0
    modified_runs = 0

    for paragraph in iter_paragraphs(doc):
        for run in paragraph.runs:
            total_runs += 1

            style_name = getattr(run.style, "name", None)

            if restrict_style and style_name != restrict_style:
                continue

            original_text = run.text
            if not original_text:
                continue

            new_text = transliterate_run_text(original_text)
            if new_text != original_text:
                modified_runs += 1
                LOGGER.debug("Run modified: %r -> %r", original_text, new_text)

                apply_tamil_font = bool(
                    tamil_font and should_apply_tamil_font(original_text)
                )

                if not dry_run:
                    # Only update text; leave all formatting intact.
                    run.text = new_text
                    if apply_tamil_font:
                        # Set the logical font name via python-docx.
                        run.font.name = tamil_font

                        # Additionally override the complex-script font (w:cs)
                        # in the underlying XML so that Word uses the Tamil
                        # font for Indic text rather than the old Sanskrit
                        # complex-script font.
                        try:
                            r = run._element  # type: ignore[attr-defined]
                            r_pr = r.get_or_add_rPr()  # type: ignore[call-arg]
                            r_fonts = getattr(r_pr, "rFonts", None)
                            if r_fonts is None:
                                r_fonts = OxmlElement("w:rFonts")
                                r_pr.append(r_fonts)

                            # Set all font families explicitly, including cs.
                            r_fonts.set(qn("w:ascii"), tamil_font)
                            r_fonts.set(qn("w:hAnsi"), tamil_font)
                            r_fonts.set(qn("w:cs"), tamil_font)
                            r_fonts.set(qn("w:hint"), "cs")
                        except Exception:
                            LOGGER.debug(
                                "Unable to update complex-script font for run.",
                                exc_info=True,
                            )

    LOGGER.info("Processed %d runs, modified %d runs", total_runs, modified_runs)
    return total_runs, modified_runs


def build_arg_parser() -> argparse.ArgumentParser:
    """Create and return the argument parser for the CLI."""
    parser = argparse.ArgumentParser(
        description=(
            "Transliterate Sanskrit (Devanagari) text in a DOCX file to Tamil "
            "script while leaving all other text and formatting unchanged."
        )
    )
    parser.add_argument(
        "input",
        metavar="INPUT",
        help="Path to the input .docx file.",
    )
    parser.add_argument(
        "-o",
        "--output",
        metavar="OUTPUT",
        help="Path for the output .docx file (default: <input_stem>_ta.docx).",
    )
    parser.add_argument(
        "--style",
        metavar="STYLE_NAME",
        help=(
            "Optional: only transliterate runs whose style name matches this "
            "value (e.g. 'Sanskrit')."
        ),
    )
    parser.add_argument(
        "--tamil-font",
        metavar="FONT_NAME",
        default="Vijaya",
        help=(
            "Font name to apply to transliterated Tamil text for runs that were "
            "entirely Sanskrit before transliteration (default: Vijaya)."
        ),
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help=(
            "Analyze and log which runs would be modified, but do not write any "
            "changes to disk."
        ),
    )
    parser.add_argument(
        "-v",
        "--verbose",
        action="store_true",
        help="Enable verbose (DEBUG-level) logging output.",
    )
    return parser


def configure_logging(verbose: bool) -> None:
    """Configure the root logger based on the verbose flag."""
    level = logging.DEBUG if verbose else logging.INFO
    logging.basicConfig(level=level, format="%(levelname)s: %(message)s")


def resolve_output_path(input_path: Path, output_arg: str | None) -> Path:
    """
    Determine the output path from the input path and optional CLI override.

    Default pattern: <input_stem>_ta.docx in the same directory.
    """
    if output_arg:
        return Path(output_arg)

    stem = input_path.stem
    suffix = input_path.suffix or ".docx"
    return input_path.with_name(f"{stem}_ta{suffix}")


def main(argv: Sequence[str] | None = None) -> int:
    parser = build_arg_parser()
    args = parser.parse_args(argv)

    configure_logging(args.verbose)

    if Document is None:
        LOGGER.error(
            "Missing dependency 'python-docx'. Install it with: pip install python-docx"
        )
        return 1

    input_path = Path(args.input)
    if not input_path.is_file():
        LOGGER.error("Input file not found: %s", input_path)
        return 1

    if input_path.suffix.lower() != ".docx":
        LOGGER.warning(
            "Input file does not have a .docx extension: %s (continuing anyway)",  # noqa: E501
            input_path,
        )

    output_path = resolve_output_path(input_path, args.output)

    if not args.dry_run:
        if output_path.resolve() == input_path.resolve():
            LOGGER.error(
                "Output path must be different from input path when not using --dry-run."  # noqa: E501
            )
            return 1

        output_dir = output_path.parent
        output_dir.mkdir(parents=True, exist_ok=True)

        LOGGER.info("Copying input DOCX to output path: %s -> %s", input_path, output_path)  # noqa: E501
        shutil.copy2(input_path, output_path)
        doc_path = output_path
    else:
        # In dry-run mode, we do not write anything; open the input directly.
        LOGGER.info("Running in dry-run mode; no files will be modified.")
        doc_path = input_path

    LOGGER.info("Loading document: %s", doc_path)
    doc = Document(str(doc_path))

    total_runs, modified_runs = process_document(
        doc,
        dry_run=args.dry_run,
        restrict_style=args.style,
        tamil_font=args.tamil_font,
    )

    if not args.dry_run:
        LOGGER.info("Saving transliterated document to: %s", output_path)
        doc.save(str(output_path))
        # Optional basic validation: try reopening the file.
        try:
            _ = Document(str(output_path))
            LOGGER.info("Successfully saved and re-opened output document.")
        except Exception as exc:  # pragma: no cover - defensive
            LOGGER.warning("Output document saved but could not be re-opened: %s", exc)
    else:
        LOGGER.info(
            "Dry-run complete. %d of %d runs would be modified.", total_runs, modified_runs  # noqa: E501
        )

    # Non-zero exit code if there was nothing to do could be considered,
    # but for now we always return 0 unless there was an earlier error.
    return 0


if __name__ == "__main__":  # pragma: no cover
    raise SystemExit(main())

